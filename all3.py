import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO
import PyPDF2

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(page_title="AA Guru", layout="wide", page_icon="🎓")


# --- SISIPAN KODE CSS & HTML UNTUK TAMPILAN PROFESIONAL ---
# --- SISIPAN KODE CSS & HTML UNTUK TAMPILAN PROFESIONAL ---
# --- SISIPAN KODE CSS & HTML UNTUK TAMPILAN PROFESIONAL (DESAIN BARU) ---
def apply_custom_ui():
    st.markdown(
        """
        <style>
        /* Impor Font Modern */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
        
        /* Pengaturan Global */
        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif;
            color: #374151; /* Warna teks abu-abu gelap agar tidak terlalu kontras */
        }

        /* Kontainer Hasil AI */
        .ai-output-card {
            background-color: #ffffff;
            padding: 2.5rem;
            margin-top: 1.5rem;
            line-height: 1.7;
            font-size: 1.05rem;
        }

        /* 1. JUDUL UTAMA (Garis Putus-putus Hijau) */
        .ai-output-card h1 {
            border: 3px dotted #84CC16; /* Hijau lime */
            padding: 15px;
            text-align: center;
            color: #1F2937;
            font-weight: 700;
            text-transform: uppercase;
            font-size: 1.5rem;
            margin-bottom: 2rem;
            border-radius: 4px;
        }

        /* 2. BAGIAN UTAMA (Kotak Biru Muda - A. IDENTITAS) */
        .ai-output-card h2 {
            background-color: #F0F9FF; /* Biru sangat muda */
            border: 1px solid #BAE6FD; /* Garis tepi biru */
            color: #1E3A8A;
            padding: 12px 20px;
            border-radius: 8px; /* Sudut melengkung */
            font-size: 1.15rem;
            font-weight: 700;
            margin-top: 2rem !important;
            margin-bottom: 1.5rem !important;
        }

        
        /* 3. SUB-BAGIAN (Label Pil Ungu Muda - Paksa dengan !important) */
        .ai-output-card h3 {
            background-color: #EEF2FF !important; 
            color: #312E81 !important;
            padding: 8px 20px !important;
            border-radius: 25px !important; /* Lengkungan pil yang lebih bulat */
            display: inline-block !important; /* Agar background membungkus teks saja */
            font-size: 1.05rem !important;
            font-weight: 700 !important;
            border-bottom: none !important;
            margin-top: 1.5rem !important;
            margin-bottom: 1rem !important;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1) !important; /* Sedikit bayangan agar menonjol */
        }
        
        /* 4. TABEL (Identitas dengan Kolom Kiri Hijau) */
        .ai-output-card table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 1.5rem;
        }
        /* Menyembunyikan Header Markdown (opsional agar mirip gambar) */
        .ai-output-card thead {
            display: none; 
        }
        .ai-output-card td {
            padding: 12px 15px;
            border-bottom: 1px solid #E5E7EB; /* Garis bawah saja yang terlihat */
            border-top: none;
            border-left: none;
            border-right: none;
        }
        /* Kolom pertama warna hijau muda */
        .ai-output-card td:first-child {
            background-color: #F7FEE7; /* Hijau pastel sangat muda */
            font-weight: 600;
            width: 35%;
            color: #4B5563;
        }
        </style>
    """,
        unsafe_allow_html=True,
    )


# Panggil Fungsi UI
apply_custom_ui()

# --- FUNGSI PARSER DOCX ---
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO


# --- 🪄 FUNGSI SIHIR: MENAMBAHKAN WARNA LATAR (BACKGROUND) DI WORD ---
# --- 🪄 FUNGSI SIHIR: WARNA LATAR PARAGRAF & SEL TABEL DI WORD ---
def set_paragraph_bg_color(paragraph, color_hex):
    """Fungsi khusus memanipulasi XML Word untuk memberi warna latar paragraf"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex) 
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(shading_elm)

def set_cell_bg_color(cell, color_hex):
    """Fungsi khusus memanipulasi XML Word untuk memberi warna latar pada SEL TABEL"""
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    tcPr.append(shading_elm)

# --- FUNGSI PARSER DOCX (VERSI ULTIMATE TAMPILAN WEB) ---
def add_markdown_paragraph(doc_or_cell, text, style=None):
    if style:
        p = doc_or_cell.add_paragraph(style=style)
    else:
        p = doc_or_cell.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 != 0: 
            run.bold = True
    return p

def create_formatted_docx(text, title):
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    doc = Document()
    
    title_head = doc.add_heading(title, 0)
    title_head.alignment = 1 
    
    lines = text.split('\n')
    is_table = False
    table_data = []
    in_identitas_block = False

    def render_table(is_identitas=False):
        if not table_data: return
        valid_data = [row for row in table_data if not all(c.strip() == '-' or c.strip() == '' for c in row[0])]
        if not valid_data: return
        
        cols_count = max(len(row) for row in valid_data)
        table = doc.add_table(rows=len(valid_data), cols=cols_count)
        
        # Tabel identitas pakai garis tipis bawaan, tabel biasa garis tegas
        table.style = 'Table Grid' if not is_identitas else 'Normal Table'
            
        for i, row in enumerate(valid_data):
            for j, cell_text in enumerate(row):
                if j < cols_count:
                    clean_cell = cell_text.replace('**', '').strip()
                    table.cell(i, j).text = clean_cell
                    
                    # 💡 PEWARNAAN TABEL IDENTITAS (Warna Hijau Muda di Kolom Kiri)
                    if is_identitas and j == 0:
                        set_cell_bg_color(table.cell(i, j), "F7FEE7")
                    
                    if not is_identitas and i == 0:
                        for paragraph in table.cell(i, j).paragraphs:
                            for run in paragraph.runs: run.bold = True

    for line in lines:
        clean_line = line.strip()
        leading_spaces = len(line) - len(line.lstrip()) 
        
        if "## A. IDENTITAS" in clean_line.upper() or "## A. IDENTITAS MODUL" in clean_line.upper():
            if is_table:
                render_table()
                table_data = []
                is_table = False
            in_identitas_block = True
            heading_text = clean_line.replace('##', '').replace('**', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            set_paragraph_bg_color(p, "EFF6FF")
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            continue

        # 💡 PENANGKAP SUB-JUDUL (Heading 3) - Tampilan PIL UNGU
        elif clean_line.startswith('### '):
            if is_table:
                render_table(is_identitas=in_identitas_block)
                table_data = []
                is_table = False
                in_identitas_block = False
                
            heading_text = clean_line.replace('###', '').replace('**', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x31, 0x2E, 0x81) # Teks Ungu Tua
            
            set_paragraph_bg_color(p, "EEF2FF") # Latar Ungu Muda
            
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Pt(18) # Sedikit menjorok ke dalam
            continue

        elif clean_line.startswith('## '):
            if is_table:
                render_table(is_identitas=in_identitas_block)
                table_data = []
                is_table = False
            in_identitas_block = False
            heading_text = clean_line.replace('##', '').replace('**', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            set_paragraph_bg_color(p, "EFF6FF")
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            continue

        if clean_line.startswith('|') and clean_line.endswith('|'):
            if '---' in clean_line: continue
            cells = [c.strip() for c in clean_line.strip('|').split('|')]
            table_data.append(cells)
            is_table = True
            continue
        else:
            if is_table and clean_line == "":
                continue 
            elif is_table and not clean_line.startswith('|'):
                render_table(is_identitas=in_identitas_block)
                table_data = []
                is_table = False
                in_identitas_block = False
        
        if not clean_line: continue
            
        if clean_line.startswith('# '):
            heading_text = clean_line.replace('#', '').replace('**', '').strip()
            doc.add_heading(heading_text, level=1)
            
        elif re.match(r'^[a-z]\.\s', clean_line):
            p = add_markdown_paragraph(doc, clean_line)
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.first_line_indent = Pt(-18)

        elif re.match(r'^\d+\.\s', clean_line):
            p = add_markdown_paragraph(doc, clean_line)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)
            
        elif clean_line.startswith(('* ', '- ')):
            text_part = clean_line[2:]
            if leading_spaces >= 2:
                p = add_markdown_paragraph(doc, "○   " + text_part)
                p.paragraph_format.left_indent = Pt(54)
                p.paragraph_format.first_line_indent = Pt(-18)
            else:
                p = add_markdown_paragraph(doc, "•   " + text_part)
                p.paragraph_format.left_indent = Pt(36)
                p.paragraph_format.first_line_indent = Pt(-18)
        
        else:
            add_markdown_paragraph(doc, clean_line)
            
    if is_table:
        render_table(is_identitas=in_identitas_block)
        
    target_stream = BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

# --- FUNGSI GLOBAL ---
def call_gemini_ai(api_key, prompt):
    try:
        genai.configure(api_key=api_key)
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        available_models = [
            m.name
            for m in genai.list_models()
            if "generateContent" in m.supported_generation_methods
        ]
        selected_model = next(
            (m for m in available_models if "1.5-flash" in m), available_models[0]
        )
        model = genai.GenerativeModel(
            model_name=selected_model, safety_settings=safety_settings
        )
        response = model.generate_content(prompt)
        return (
            response.text
            if response.candidates and response.candidates[0].content.parts
            else "ERROR: Respons kosong."
        )
    except Exception as e:
        return f"ERROR: {str(e)}"


def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    return "".join([page.extract_text() for page in pdf_reader.pages])


# --- 2. SESSION STATE ---
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "tp_result" not in st.session_state:
    st.session_state.tp_result = ""
if "atp_result" not in st.session_state:
    st.session_state.atp_result = ""
if "modul_result" not in st.session_state:
    st.session_state.modul_result = ""
if "page_modul" not in st.session_state:
    st.session_state.page_modul = 1
if "data_modul" not in st.session_state:
    st.session_state.data_modul = {}
if "fase_terpilih" not in st.session_state:
    st.session_state.fase_terpilih = "Fase A"
if "soal_result" not in st.session_state:
    st.session_state.soal_result = None
if "kisikisi_result" not in st.session_state:
    st.session_state.kisikisi_result = None
if "list_topik" not in st.session_state:
    st.session_state.list_topik = [{"nama": "", "jumlah": 5}]

# --- 3. SIDEBAR NAVIGATION ---
page = st.sidebar.radio(
    "Tahapan Kerja:",
    [
        "1. Bedah CP & TP",
        "2. Alur (ATP) & Pemetaan JP",
        "3. Modul Ajar Expert",
        "4. Generator Soal & Kisi-kisi",
    ],
)
st.sidebar.divider()
st.sidebar.markdown("### 🔑 Akses")
st.session_state.api_key = st.sidebar.text_input(
    "API Key Gemini:", type="password", value=st.session_state.api_key
)

# --- 4. LOGIKA HALAMAN ---

if page == "1. Bedah CP & TP":
    st.header("📋 Tahap 1: Bedah CP & Tujuan Pembelajaran (TP)")
    uploaded_file = st.file_uploader("Unggah PDF CP (Opsional):", type="pdf")
    initial_cp = read_pdf(uploaded_file) if uploaded_file else ""

    col1, col2 = st.columns([3, 1])
    with col1:
        cp_input = st.text_area(
            "Tempel Teks CP BSKAP 046/2025:", value=initial_cp, height=250
        )
    with col2:
        st.session_state.fase_terpilih = st.selectbox(
            "Pilih Fase:", ["Fase A", "Fase B", "Fase C", "Fase D", "Fase E", "Fase F"]
        )

    if st.button("Generate Analisis & TP", type="primary", use_container_width=True):
        if not st.session_state.api_key or not cp_input:
            st.warning("Mohon lengkapi API Key dan teks CP.")
        else:
            with st.spinner("AI sedang membedah CP..."):
                prompt = f"""Bertindaklah sebagai ahli kurikulum Spesialis Kurikulum Kemendikbudristek (Update BSKAP 046/2025). Analisis CP berikut: {cp_input}. 
                1. Buat tabel dengan format yang rapi analisis Kompetensi & Materi Pokok. 
                2. Turunkan menjadi TP yang dibagi otomatis per kelas dalam {st.session_state.fase_terpilih} secara scaffolding.
                    Instruksi Analisis:
                    a. Dekonstruksi: Pisahkan identifikasi Kompetensi Kata Kerja Operasional(KKO) dan Konten (Materi Esensial).
                    b. Perumusan TP: Buat Tujuan Pembelajaran yang konkret, terukur, dan mencakup aspek pemahaman.
                    c. Penyusunan ATP: Urutkan TP secara logis dan prasyarat sesuai prinsip Panduan Pembelajaran dan Asesmen (PPA) Kurikulum Merdeka 2025/2026.
                    d. Deep Learning Integration: Berikan saran aktivitas belajar yang:
                        - Mindful: Membangun kesadaran diri siswa akan tujuan belajar.
                        - Meaningful: Menghubungkan konteks nyata/masalah otentik.
                        - Joyful: Menantang namun menyenangkan (Flow state).
                    e. Output harus dalam tabel Markdown yang rapi."""

                st.session_state.tp_result = call_gemini_ai(
                    st.session_state.api_key, prompt
                )
                st.rerun()

    if st.session_state.tp_result:
        st.markdown(
            f'<div class="ai-output-card">{st.session_state.tp_result}</div>',
            unsafe_allow_html=True,
        )
        st.download_button(
            "📥 Unduh TP (Docx)",
            create_formatted_docx(st.session_state.tp_result, "Analisis CP dan TP"),
            "TP_Analisis.docx",
        )

elif page == "2. Alur (ATP) & Pemetaan JP":
    st.header("🗺️ Tahap 2: Alur Tujuan Pembelajaran (ATP) & JP")
    if not st.session_state.tp_result:
        st.error("⚠️ Selesaikan Tahap 1 terlebih dahulu.")
    else:
        if st.button(
            "Generate Tabel ATP & Pemetaan JP", type="primary", use_container_width=True
        ):
            with st.spinner("AI sedang menyusun alur..."):
                prompt = f"""Buatlah tabel ATP berdasarkan data TP ini: {st.session_state.tp_result}.
                Buatlah tabel Alur Tujuan Pembelajaran (ATP) dengan Output harus dalam tabel Markdown yang rapi untuk {st.session_state.fase_terpilih}.
                buatkan secara lengkap dan rinci, WAJIB tampilkan pada kolom: No, Capaian Pembelajaran (CP), Elemen, Kelas, Semester, TP, Materi Pokok, Alokasi Waktu (JP)."""
                st.session_state.atp_result = call_gemini_ai(
                    st.session_state.api_key, prompt
                )
                st.rerun()

        if st.session_state.atp_result:
            st.markdown(
                f'<div class="ai-output-card">{st.session_state.atp_result}</div>',
                unsafe_allow_html=True,
            )
            st.download_button(
                "📥 Unduh ATP (Docx)",
                create_formatted_docx(
                    st.session_state.atp_result, "Alur Tujuan Pembelajaran"
                ),
                "ATP_JP.docx",
            )

elif page == "3. Modul Ajar Expert":
    d = st.session_state.data_modul

    if st.session_state.page_modul == 1:
        st.title("📝 Penyusunan Modul Ajar")
        st.progress(0.33)

        with st.form("form_input_modul"):
            col1, col2 = st.columns(2)
            with col1:
                nama = st.text_input(
                    "Nama Guru",
                    value=d.get("nama", ""),
                    placeholder="Contoh: Iman Nuriman, ST.",
                )
                unit = st.text_input(
                    "Unit Kerja",
                    value=d.get("unit", ""),
                    placeholder="Contoh: SMP Negeri 1 Pangalengan",
                )
                mapel = st.text_input("Mata Pelajaran", value=d.get("mapel", ""))
                fase_input = st.selectbox(
                    "Fase", ["A", "B", "C", "D", "E", "F"], index=0
                )

            with col2:
                kelas = st.text_input("Kelas", value=d.get("kelas", ""))
                semester = st.selectbox("Semester", ["1 (Ganjil)", "2 (Genap)"])
                jp = st.text_input(
                    "Alokasi Waktu",
                    value=d.get("jp", ""),
                    placeholder="Contoh: 2 x 40 Menit",
                )
                topik = st.text_input("Topik Pembelajaran", value=d.get("topik", ""))

            st.markdown("#### 🎯 Dimensi Profil Lulusan (DPL)")
            dimensi_dpl = st.multiselect(
                "Pilih Dimensi:",
                [
                    "Keimanan",
                    "Kewargaan",
                    "Penalaran Kritis",
                    "Kreativitas",
                    "Kolaborasi",
                    "Kemandirian",
                    "Kesehatan",
                    "Komunikasi",
                ],
                default=["Penalaran Kritis"],
            )

            st.markdown("#### ⚙️ Metode Pembelajaran")
            model_belajar = st.selectbox(
                "Model Pembelajaran",
                [
                    "PBL",
                    "PjBL",
                    "Inquiry",
                    "Cooperative",
                    "Discovery",
                    "Berdiferensiasi",
                ],
            )
            pertemuan = st.number_input("Jumlah Pertemuan", min_value=1, value=1)
            kondisi_khusus = st.text_area(
                "Instruksi Tambahan:", value=d.get("kondisi_khusus", "")
            )

            submit = st.form_submit_button(
                "Lanjut ke Konfirmasi →", use_container_width=True
            )
            if submit:
                st.session_state.data_modul = {
                    "nama": nama,
                    "unit": unit,
                    "mapel": mapel,
                    "fase": fase_input,
                    "kelas": kelas,
                    "semester": semester,
                    "jp": jp,
                    "pertemuan": pertemuan,
                    "topik": topik,
                    "model": model_belajar,
                    "kondisi_khusus": kondisi_khusus,
                    "dimensi_dpl": dimensi_dpl,
                }
                st.session_state.page_modul = 2
                st.rerun()

    elif st.session_state.page_modul == 2:
        st.title("🔍 Konfirmasi Kerangka Pembelajaran")
        data = st.session_state.data_modul
        with st.container(border=True):
            c1, c2 = st.columns(2)
            with c1:
                st.write(f"**Nama:** {data.get('nama')}")
                st.write(f"**Mapel:** {data.get('mapel')}")
            with c2:
                st.write(f"**Topik:** {data.get('topik')}")
                st.write(f"**Alokasi:** {data.get('jp')}")

        col_bt1, col_bt2 = st.columns(2)
        if col_bt1.button("⬅️ Edit Kembali", use_container_width=True):
            st.session_state.page_modul = 1
            st.rerun()
        if col_bt2.button("🚀 GENERATE", type="primary", use_container_width=True):
            st.session_state.page_modul = 3
            st.rerun()  # [PERBAIKAN] Menambahkan tanda kurung

    elif st.session_state.page_modul == 3:
        st.title("✨ Hasil Modul Ajar")
        d = st.session_state.data_modul

        # [PERBAIKAN] Cek apakah hasil sudah ada, jika belum baru panggil AI
        if not st.session_state.modul_result:
            prompt = f"""Bertindaklah sebagai Guru Ahli Kurikulum 2026 terbaru. Buatlah **Modul Ajar** lengkap dengan pendekatan **Deep Learning** (Mindful, Meaningful, Joyful).
                        IDENTITAS: 
                        Nama: {d['nama']}, Unit: {d['unit']}, Mapel: {d['mapel']}, Fase/Kelas: {d['fase']}/{d['kelas']}, Semester: {d['semester']}, Alokasi: {d['jp']}, Topik: {d['topik']}.

                        INSTRUKSI KHUSUS PEMBAGIAN ALOKASI WAKTU (WAJIB DIPATUHI): jika Fase {d['fase']} = A atau B atau C maka setiap 1 jam pelajaran (JP) = 35 Menit, jika Fase {d['fase']} = D maka setiap 1 jam pelajaran (JP) = 40 Menit, jika Fase {d['fase']} = E atau F maka setiap 1 jam pelajaran (JP) = 45 Menit.

                        INSTRUKSI KHUSUS DARI GURU (WAJIB DIINTEGRASIKAN):
                        {d['kondisi_khusus'] if d['kondisi_khusus'] else "Tidak ada instruksi tambahan."}
                    
                        ATURAN FORMAT PENULISAN MARKDOWN (WAJIB DIPATUHI UNTUK TAMPILAN WEB):
                        1. Judul Utama: Gunakan format Heading 1. Contoh: # MODUL AJAR: {d['topik'].upper()}
                        2. Bagian Utama (A, B, C, dst): WAJIB gunakan format Heading 2 (`## `). Contoh: ## A. IDENTITAS MODUL
                        3. Sub-bagian: WAJIB gunakan format Heading 3 (`### `) agar tampil sebagai tombol/pil. Contoh: ### Tujuan Pembelajaran, ### Praktik Pedagogis, dll.
                        4. Khusus bagian A. IDENTITAS MODUL: Buat dalam bentuk tabel Markdown 2 kolom TANPA baris header judul. Kolom 1 untuk atribut (contoh: Mata Pelajaran), Kolom 2 untuk isinya.

                        STRUKTUR MODUL YANG WAJIB DIIKUTI:
                        # MODUL AJAR: {d['topik'].upper()}                
                            ## A. IDENTITAS MODUL
                                (Buat tabel identitas 2 kolom di sini sesuai aturan)
                            ## B. CAPAIAN PEMBELAJARAN (CP): Jabarkan elemen dan rumusan CP sesuai topik {d['topik']} yang mengacu pada BSKAP Nomor 046/H/KR/2025.
                            ## C. DIMENSI PROFIL LULUSAN (DPL)
                                    ### Dimensi Profil Lulusan 
                                    Integrasikan dimensi {', '.join(d['dimensi_dpl'])} secara eksplisit dalam aktivitas.
                            ## D. CAKUPAN MATERI: rumusan ruang lingkup materi apa saja yang akan dilaksanakan dalam pembelajaran sesuai dengan topik {d['topik']}.        
                            ## D. DESAIN PEMBELAJARAN
                                    ### TUJUAN PEMBELAJARAN:Susun Tujuan Pembelajaran (TP) dalam bentuk kalimat paragraf utuh yang sudah mengandung unsur ABCD (Audience, Behavior, Condition, Degree) di dalam kalimatnya. JANGAN membedah atau menuliskan singkatan A, B, C, D tersebut secara eksplisit. Jika rumusan TP lebih dari satu, sajikan dalam bentuk daftar (bullet/numbering)
                                    ### PRAKTIK PAEDAGOGIS :Tuliskan model {d['model']} yang dipilih untuk mencapai tujuan pembelajaran dan tuliskan sintaksnya.
                                    ### KEMITRAAN PEMBELAJARAN (OPSIONAL) :Tuliskan kegiatan kemitraan atau kolaborasi dalam dan/atau ruang lingkup sekolah, seperti: kemitraan antar guru, lintas mata pelajaran, antar murid antar kelas, antar guru lintas sekolah, orang tua, komunitas, tokoh masyarakat, dunia usaha dan dunia industri kerja, institusi, atau mitra profesional.
                                    ### LINGKUNGAN PEMBELAJARAN : Tuliskan lingkungan pembelajaran yang diinginkan dalam pembelajaran dalam budaya belajar, ruang fisik dan/atau ruang virtual agar tecipta iklim belajar yang aman, nyaman, dan saling memuliakan, contoh : memberikan kepada siswa untuk menyampaikan pendapatnya dalam ruang kelas dan dan forum diskusi pada platform daring (ruang virtual bersifat opsional).
                                    ### PEMANFAATAN DIGITAL (OPSIONAL):Tuliskan pemanfaatan digital untuk menciptakan pembelajaran yang inteaktif, kolaboratif dan kontekstual, contoh : video pembelajaran, platform pembelajaran, perpustakaan digital, forum diskusi daring, aplikasi penilaian, dan sebagainya.
                            ## F. PEMAHAMAN BERMAKNA & PERTANYAAN PEMANTIK: Berisi 3 Pertanyaan HOTS.
                            ## G. LANGKAH-LANGKAH PEMBELAJARAN (Sintaks {d['model']} buat dalam {d['pertemuan']} pertemuan): Wajib mencakup 3 Kategori Deep Learning: 1. MEMAHAMI (Berkesadaran & Bermakna), 2. MENGAPLIKASI (Berkesadaran, Bermakna, Menyenangkan), 3. MEREFLEKSI (Berkesadaran & Bermakna)
                                    ### Pendahuluan: Membangun koneksi emosional dan kesadaran (Mindful). 
                                    ### Kegiatan Inti : Eksplorasi mendalam menggunakan sintaks {d['model']}'
                                    ### Penutup: Refleksi metakognisi (Apa yang sekarang saya tahu yang sebelumnya saya tidak tahu?).
                            ## H. ASESMEN: 
                            WAJIB: Sajikan bagian asesmen dalam TABEL TERPISAH dengan ketentuan sebagai berikut:
                                    ### INSTRUMEN ASESMEN : Tuliskan instrumen asesment yang akan dipergunakan selama proses pembelajaran berlangsung dai awal sampai akhir Sajikan dalam tabel.
                                    ### TEKNIK ASESMEN: Tuliskan teknik asesment yang akan dipergunakan selama proses pembelajaran berlangsung dai awal  sampai  akhir,  apakah  menggunakan  tehnik  tes,  yaitu  :  tes  tulis,  tes  lisan,  atau  tes perbuatan dan non tes, yaitu : penilaian sejawat, penilaian diri, penilaian produk, observasi, portofolio, penilaian berbasis kelas, penilaian kinerja, skala sikap, wawancara, atau sosiometri, beserta contohnya dan sajikan dalam tabel.               
                            ## I. MEDIA, ALAT, DAN SUMBER BELAJAR : 
                                    ### MEDIA DAN ALAT PEMBELAJARAN : Tuliskan media dan alat pembelajaran yang akan dipergunakan pada saat pembelajaran berlangsung untuk membantu dan/atau mempermudah pemahaman murid dalam menerima materi pembelajaran.
                                    ### SUMBER BELAJAR : Tuliskan referensi baik berupa buku, jurnal, kamus, surat kabar, majalah, website, dan/atau yang lainnya yang akan  dipakai  selama proses  pembelajaran  dalam mendukung  ketecapaian kompetensi seperti yang telah dirumuskan dalam tujuan pembelajaran di atas. Contoh penulisan referensi berupa buku dalam sumber belajar, yaitu : Haris, Mohamad, 2020, Mudah Belajar Matematika, hal. 27-32, edisi kedua, cetakan kesatu, Surabaya, Pelita Bangsa.

                            ## J. LAMPIRAN          
                                    ### LEMBAR KERJA PESERTA DIDIK (LKPD): 
                                    ### Lembar Kerja Peserta Didik (LKPD) : buatkan LKPD sesuai dengan jumlah pertemuan dan Buat instruksi tugas yang jelas, membuat siswa aktif dan mendalam setiap LKPD nya disertai RUBRIK PENILAIAN LKPD nya.
                                    ### Ringkasan Materi Mendalam
                                    ### Glosarium.

            Gunakan bahasa Indonesia yang formal namun mudah dipahami guru dan Bahasa natural tidak seperti Bahasa mesin.
            """

            with st.status(
                "🚀 AI sedang menyusun perangkat ajar...", expanded=True
            ) as status:
                # [PERBAIKAN] Simpan ke session_state agar tidak hilang
                st.session_state.modul_result = call_gemini_ai(
                    st.session_state.api_key, prompt
                )
                status.update(label="Selesai!", state="complete")

        # [PERBAIKAN] Tampilkan hasil dari session_state
        if st.session_state.modul_result:
            st.markdown(
                f'<div class="ai-output-card">{st.session_state.modul_result}</div>',
                unsafe_allow_html=True,
            )

            docx_bytes = create_formatted_docx(
                st.session_state.modul_result, f"Modul Ajar - {d['topik']}"
            )

            st.divider()
            c_dl, c_new = st.columns([3, 1])
            with c_dl:
                st.download_button(
                    label="📥 Download Modul Ajar (.docx)",
                    data=docx_bytes,
                    file_name=f"Modul_Ajar_{d['topik'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with c_new:
                if st.button("🔄 Buat Baru", use_container_width=True):
                    st.session_state.modul_result = ""
                    st.session_state.page_modul = 1
                    st.rerun()

elif page == "4. Generator Soal & Kisi-kisi":
    st.header("❓ Tahap 4: Bank Soal & Kisi-kisi")

    # Ambil data dari Modul Ajar secara otomatis (Sinkronisasi)
    d_modul = st.session_state.data_modul

    col_config1, col_config2, col_config3 = st.columns(3)
    with col_config1:
        jenjang = st.selectbox("Jenjang", ["SD", "SMP", "SMA", "SMK"], index=1)
    with col_config2:
        kelas_soal = st.text_input("Kelas", value=d_modul.get("kelas", ""))
    with col_config3:
        mapel_soal = st.text_input("Mata Pelajaran", value=d_modul.get("mapel", ""))

    st.subheader("📚 Manajemen Topik")
    # Validasi list_topik agar tidak error saat iterasi
    if not st.session_state.list_topik:
        st.session_state.list_topik = [{"nama": "", "jumlah": 5}]

    for i, item in enumerate(st.session_state.list_topik):
        c1, c2, c3 = st.columns([3, 1, 0.5])
        default_topik = (
            item["nama"]
            if item["nama"]
            else (d_modul.get("topik", "") if i == 0 else "")
        )
        st.session_state.list_topik[i]["nama"] = c1.text_input(
            f"Topik {i+1}", value=default_topik, key=f"topik_input_{i}"
        )
        st.session_state.list_topik[i]["jumlah"] = c2.number_input(
            f"Jml Soal", min_value=1, value=item["jumlah"], key=f"jml_input_{i}"
        )
        if c3.button("🗑️", key=f"del_topik_{i}"):
            if len(st.session_state.list_topik) > 1:
                st.session_state.list_topik.pop(i)
                st.rerun()

    if st.button("➕ Tambah Topik Baru"):
        st.session_state.list_topik.append({"nama": "", "jumlah": 5})
        st.rerun()

    # Form Pengaturan Soal
    with st.form("form_soal_expert"):
        st.write("### ⚙️ Pengaturan Jenis & Jumlah")
        f1, f2, f3 = st.columns(3)
        n_pg = f1.number_input("Jumlah PG", min_value=0, value=10)
        n_essay = f2.number_input("Jumlah Essay", min_value=0, value=5)
        n_bs = f3.number_input("Jumlah B/S", min_value=0, value=0)

        st.write("### 📊 Tingkat Kesulitan (%)")
        diff_col1, diff_col2, diff_col3 = st.columns(3)
        p_mudah = diff_col1.number_input("Mudah (C1-C2) %", value=30)
        p_sedang = diff_col2.number_input("Sedang (C3-C4) %", value=50)
        p_sulit = diff_col3.number_input("Sulit (C5-C6) %", value=20)

        st.write("### 🎨 Pengaturan Gambar")
        img_c1, img_c2 = st.columns(2)
        cb_gambar = img_c1.checkbox("Sertakan Prompt Gambar Detail", value=True)
        n_gambar = img_c1.number_input(
            "Jumlah Soal Stimulus Gambar", min_value=0, value=2
        )
        gaya_gambar = img_c2.selectbox(
            "Gaya Visual",
            [
                "Diagram Teknis",
                "Ilustrasi Edukasi",
                "Foto Realistik",
                "Sketsa",
                "Gambar style kartun 3d",
            ],
        )

        generate_btn = st.form_submit_button(
            "🚀 Generate Bank Soal", use_container_width=True
        )

        if generate_btn:
            if not st.session_state.api_key:
                st.error("Masukkan API Key di Sidebar!")
            else:
                with st.spinner("AI sedang merancang soal berkualitas..."):
                    # Menyiapkan rincian topik untuk prompt
                    valid_topik = [
                        t
                        for t in st.session_state.list_topik
                        if t["nama"].strip() != ""
                    ]
                    rincian_str = "\n".join(
                        [f"- {t['nama']}: {t['jumlah']} soal" for t in valid_topik]
                    )

                    prompt_visual = ""
                    if cb_gambar and n_gambar > 0:
                        prompt_visual = f"\nSertakan {n_gambar} soal dengan [Gambar: Prompt: <deskripsi detail>] gaya {gaya_gambar}."

                    prompt_soal = (
                        f"""Anda adalah seorang spesialis evaluasi pendidikan. Buat naskah soal {jenjang} {mapel_soal} Kelas {kelas_soal}. yang disesuaikan dengan kaidah-kaidah penyusunan soal yang baik dan benar sebagai berikut:
                        1. Kaidah Substansi/Materi (Kesesuaian):
                             a. Sesuai Indikator: Soal harus mengukur perilaku dan materi yang ditetapkan dalam kisi-kisi.
                             b. Pilihan Jawaban Homogen: Semua pilihan jawaban (pengecoh) harus logis, masuk akal, dan homogen dari segi materi.
                             c. Satu Jawaban Benar: Hanya ada satu kunci jawaban yang benar untuk setiap soal.
                             d. Tidak SARA: Soal tidak boleh menyinggung isu SARA, politik, pornografi, atau kekerasan. 
                        2. Kaidah Konstruksi (Teknis Soal):
                             a. Pokok Soal Jelas: Pokok soal (stem) dirumuskan secara jelas, tegas, dan tidak menimbulkan penafsiran ganda.
                             b. Hindari Petunjuk Jawaban: Pokok soal jangan memberi petunjuk ke arah jawaban benar.
                             c. Negatif Ganda: Hindari penggunaan pernyataan yang bersifat negatif ganda.
                             d. Panjang Pilihan Jawaban: Panjang rumusan pilihan jawaban (pilihan ganda) harus relatif sama.
                             e. Pengecoh Berfungsi: Pengecoh (distractor) harus berfungsi, logis, dan dipilih oleh peserta didik yang kurang paham materi.
                             f. Grafik/Tabel Jelas: Gambar, grafik, tabel, atau diagram harus jelas dan berfungsi dalam soal. 
                        3. Kaidah Bahasa:
                             a. Bahasa Baku: Menggunakan bahasa Indonesia yang baik dan benar (baku) sesuai ejaan (EYD).
                             b. Komunikatif: Rumusan soal harus komunikatif dan mudah dipahami sesuai jenjang pendidikan peserta didik.
                             c. Tidak Ambigu: Kalimat soal tidak menimbulkan tafsiran ganda.
                             d. Bahasa Setempat: Hindari penggunaan bahasa atau istilah yang hanya berlaku di tempat tertentu (lokal/tabu). 
                        4. Kaidah Khusus: 
                             a. Jawaban Singkat: Kalimat harus dirumuskan agar jawaban yang dihasilkan benar-benar singkat dan jelas.
                             b. Uraian (Essay): Rumusan soal menggunakan kata tanya yang menuntut uraian, seperti: "mengapa", "jelaskan", "uraikan".
                             c. Pedoman Penskoran: Soal uraian wajib disertai dengan pedoman penskoran atau kunci jawaban.\n"""
                        f"Materi: {rincian_str}\n"
                        f"Komposisi: {n_pg} PG, {n_essay} Essay, {n_bs} B/S.\n"
                        f"Target: Mudah {p_mudah}%, Sedang {p_sedang}%, Sulit {p_sulit}%.\n"
                        f"{prompt_visual}\n"
                        f"Aturan: Jika SD/SMP opsi A-D, jika SMA/SMK opsi A-E. "
                        f"Cantumkan Level Kognitif di awal soal. Sertakan Kunci Jawaban."
                    )

                    st.session_state.soal_result = call_gemini_ai(
                        st.session_state.api_key, prompt_soal
                    )
                    st.rerun()

    # [PERBAIKAN] Tampilkan hasil soal di luar form agar tombol kisi-kisi berfungsi
    if st.session_state.soal_result:
        st.markdown(
            f'<div class="ai-output-card">{st.session_state.soal_result}</div>',
            unsafe_allow_html=True,
        )
        btn_soal_docx = create_formatted_docx(
            st.session_state.soal_result, f"Bank Soal {mapel_soal}"
        )
        st.download_button(
            "📥 Unduh Bank Soal (DOCX)", btn_soal_docx, f"Soal_{mapel_soal}.docx"
        )

        st.divider()
        st.subheader("📋 Generator Kisi-kisi (BSKAP 046/2025)")
        if st.button(
            "✨ Buat Kisi-kisi Otomatis", type="primary", use_container_width=True
        ):
            with st.spinner("Memetakan soal ke CP BSKAP 046/2025..."):
                prompt_kisi = (
                    f"Buatlah TABEL kisi-kisi berdasarkan soal ini: {st.session_state.soal_result}. "
                    "Gunakan referensi BSKAP No. 046/H/KR/2025. Kolom: No, CP, Elemen, Indikator Soal, Level, Bentuk Soal."
                )
                st.session_state.kisikisi_result = call_gemini_ai(
                    st.session_state.api_key, prompt_kisi
                )
                st.rerun()

        if st.session_state.kisikisi_result:
            st.markdown(
                f'<div class="ai-output-card">{st.session_state.kisikisi_result}</div>',
                unsafe_allow_html=True,
            )
            btn_kisi_docx = create_formatted_docx(
                st.session_state.kisikisi_result, "Kisi-kisi Instrumen Penilaian"
            )
            st.download_button(
                "📥 Unduh Kisi-kisi (DOCX)", btn_kisi_docx, "Kisi_Kisi.docx"
            )
